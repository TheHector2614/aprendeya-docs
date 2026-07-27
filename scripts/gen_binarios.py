import docx, openpyxl
from pathlib import Path

BASE = Path("raw")

def heading(d, level, text):
    d.add_heading(text, level)

def para(d, text):
    d.add_paragraph(text)

sh = BASE / "sharepoint"

# ── LEG-004 Contrato de Instructores ──
d = docx.Document()
heading(d, 0, "Contrato de Prestacion de Servicios para Instructores")
heading(d, 1, "Partes")
para(d, "Entre AprendeYa SAS, identificada con NIT 901.XXX.XXX, en adelante LA PLATAFORMA, y el instructor registrado en la plataforma, en adelante EL INSTRUCTOR.")
heading(d, 1, "Objeto")
para(d, "EL INSTRUCTOR se obliga a crear, publicar y mantener contenido educativo en la plataforma AprendeYa, de acuerdo con los terminos establecidos en este contrato y en la Politica de Propiedad Intelectual.")
heading(d, 1, "Regalias")
para(d, "LA PLATAFORMA pagara a EL INSTRUCTOR el 70% de los ingresos netos generados por las ventas de sus cursos. Los pagos se realizaran mensualmente, dentro de los 10 dias habiles siguientes al cierre del ciclo de facturacion.")
heading(d, 1, "Obligaciones del Instructor")
para(d, "El INSTRUCTOR debera: (a) crear contenido original y de calidad, (b) responder preguntas de estudiantes en un maximo de 48 horas, (c) mantener actualizado el contenido del curso, (d) no infringir derechos de terceros.")
heading(d, 1, "Duracion y Terminacion")
para(d, "El contrato tiene duracion indefinida. Cualquiera de las partes puede terminarlo con 30 dias de aviso. LA PLATAFORMA puede terminar el contrato de forma inmediata si EL INSTRUCTOR viola las politicas de contenido.")
d.save(str(sh / "leg" / "LEG-004.docx"))
print("LEG-004.docx OK")

# ── OPE-004 Gestion de Incidentes de Seguridad ──
d = docx.Document()
heading(d, 0, "Procedimiento de Gestion de Incidentes de Seguridad")
heading(d, 1, "Clasificacion de Incidentes")
para(d, "Critico: afectacion a datos de usuarios o disponibilidad de la plataforma. Alto: vulnerabilidad explotable sin autenticacion. Medio: vulnerabilidad con requisitos de acceso previo. Bajo: problemas de configuracion menores o incumplimiento de politicas.")
heading(d, 1, "Flujo de Respuesta")
para(d, "1. Deteccion: monitoreo automatizado, reporte de usuario, o alerta de seguridad. 2. Triage: clasificar el incidente en menos de 30 minutos. 3. Contencion: aislar los sistemas afectados. 4. Erradicacion: eliminar la causa raiz. 5. Recuperacion: restaurar servicios desde backups limpios. 6. Post-mortem: documentar lecciones aprendidas en maximo 5 dias.")
heading(d, 1, "TIempos de Respuesta")
para(d, "Critico: contencion en < 15 min, resolucion en < 4 horas. Alto: contencion en < 1 hora, resolucion en < 24 horas. Medio: resolucion en < 7 dias. Bajo: resolucion en < 30 dias.")
heading(d, 1, "Notificaciones")
para(d, "Incidentes criticos: notificar al DPO y al director de tecnologia inmediatamente. Incidentes con datos personales: notificar a la autoridad de proteccion de datos dentro de las 72 horas siguientes, conforme a la ley aplicable.")
d.save(str(sh / "ope" / "OPE-004.docx"))
print("OPE-004.docx OK")

# ── CAL-002 Procedimiento PQRS ──
d = docx.Document()
heading(d, 0, "Procedimiento de Quejas, Reclamos y Sugerencias (PQRS)")
heading(d, 1, "Canales de Recepcion")
para(d, "Los estudiantes pueden presentar PQRS a traves de: formulario web en la plataforma, correo electronico (pqrs@aprendeya.com), linea telefonica (01-800-123-4567 opcion 3), y chat en vivo.")
heading(d, 1, "Categorias")
para(d, "Queja: insatisfaccion con la calidad del servicio. Reclamo: solicitud de cumplimiento de una obligacion. Sugerencia: propuesta de mejora. Solicitud de informacion: peticion de datos sobre la plataforma o sus servicios.")
heading(d, 1, "TIempos de Respuesta")
para(d, "Quejas y reclamos: 15 dias habiles prorrogables por 15 dias mas. Sugerencias: 10 dias habiles. Solicitudes de informacion: 10 dias habiles. Peticiones de datos personales: 15 dias habiles.")
heading(d, 1, "Escalamiento")
para(d, "Si el estudiante no esta satisfecho con la respuesta, puede solicitar revision en segunda instancia ante el Comite de Atencion al Usuario, que sesiona quincenalmente y responde en un maximo de 20 dias habiles.")
d.save(str(sh / "cal" / "CAL-002.docx"))
print("CAL-002.docx OK")

# ── RH-005 Bienestar y Salud Mental ──
d = docx.Document()
heading(d, 0, "Politica de Bienestar y Salud Mental")
heading(d, 1, "Programas de Apoyo")
para(d, "AprendeYa ofrece los siguientes programas: linea de apoyo psicologico confidencial (disponible 24/7 via telefono y chat), hasta 8 sesiones de terapia virtual por ano cubiertas por la empresa, talleres mensuales de manejo del estres y mindfulness.")
heading(d, 1, "Equilibrio Vida-Laboral")
para(d, "Jornada laboral maxima de 40 horas semanales. Derecho a la desconexion digital fuera del horario laboral. No se espera respuesta a correos o mensajes despues de las 6pm ni los fines de semana. Reuniones internas programadas unicamente dentro del horario de 9am a 5pm.")
heading(d, 1, "Dias de Bienestar")
para(d, "Cada colaborador tiene derecho a 2 dias de bienestar por semestre, que pueden utilizarse sin justificacion para atender su salud fisica o mental. Estos dias no descuentan de vacaciones ni requieren soporte medico.")
heading(d, 1, "Prevencion del Acoso")
para(d, "Se tolera cero conductas de acoso laboral, acoso sexual o discriminacion. Cualquier colaborador que experimente o presencie estas conductas puede reportarlas de forma confidencial a bienestar@aprendeya.com o a traves de la linea etica.")
d.save(str(sh / "rh" / "RH-005.docx"))
print("RH-005.docx OK")

# ── RH-006 Codigo de Etica ──
d = docx.Document()
heading(d, 0, "Codigo de Etica y Convivencia")
heading(d, 1, "Principios Rectores")
para(d, "Respeto, integridad, transparencia, responsabilidad, equidad y confidencialidad. Todos los colaboradores, instructores y directivos deben actuar conforme a estos principios en el desempeno de sus funciones.")
heading(d, 1, "Conflicto de Intereses")
para(d, "Los colaboradores deben evitar situaciones donde sus intereses personales entren en conflicto con los intereses de AprendeYa. Deben revelar cualquier conflicto potencial a su lider inmediato y abstenerse de participar en las decisiones relacionadas.")
heading(d, 1, "Uso de Recursos")
para(d, "Los recursos de la empresa (equipos, software, instalaciones, tiempo laboral) deben utilizarse unicamente para fines relacionados con el trabajo. El uso personal incidental esta permitido siempre que no afecte la productividad ni incurra en costos adicionales.")
heading(d, 1, "Regalos y Atenciones")
para(d, "No se permite aceptar regalos, atenciones o invitaciones de proveedores, clientes o terceros con un valor superior a $50 USD. Cualquier beneficio de mayor valor debe ser reportado al area de compliance.")
heading(d, 1, "Linea Etica")
para(d, "Cualquier colaborador puede reportar violaciones a este codigo de forma anonima a traves de la Linea Etica disponible en la intranet. Se prohibe cualquier tipo de represalia contra quienes reporten de buena fe.")
d.save(str(sh / "rh" / "RH-006.docx"))
print("RH-006.docx OK")

# ── COM-003 Convenios Corporativos ──
d = docx.Document()
heading(d, 0, "Programa de Convenios Corporativos")
heading(d, 1, "Beneficios por Convenio")
para(d, "Las empresas aliadas acceden a descuentos progresivos segun el numero de empleados inscritos: 10-49 empleados: 15% descuento. 50-199 empleados: 25% descuento. 200+ empleados: 35% descuento + reportes de progreso personalizados.")
heading(d, 1, "Requisitos")
para(d, "La empresa debe firmar un convenio marco con vigencia minima de 1 ano. El pago puede ser mensual, trimestral o anual. Las empresas pueden elegir entre catalogo abierto o cursos seleccionados.")
heading(d, 1, "Facturacion Corporativa")
para(d, "La facturacion se realiza a 30 dias con orden de compra. Las empresas pueden agregar o quitar empleados mensualmente. El reporte de uso se entrega el primer dia habil de cada mes.")
heading(d, 1, "Contacto Comercial")
para(d, "Las empresas interesadas pueden contactar a corporativo@aprendeya.com o al 01-800-123-4567 opcion 2. Un ejecutivo comercial asignado acompanara todo el proceso de implementacion.")
d.save(str(sh / "com" / "COM-003.docx"))
print("COM-003.docx OK")

# ── CAL-003 Politica de Satisfaccion NPS ──
d = docx.Document()
heading(d, 0, "Politica de Satisfaccion del Cliente (NPS)")
heading(d, 1, "Medicion")
para(d, "El Net Promoter Score (NPS) se mide trimestralmente a traves de la pregunta: 'En una escala del 0 al 10, que tan probable es que recomiendes AprendeYa a un amigo o colega?' Los promotores (9-10) menos los detractores (0-6) determinan el puntaje NPS.")
heading(d, 1, "Meta")
para(d, "NPS actual: 45. Meta para 2026: 55. Meta para 2027: 65. El NPS se segmenta por: tipo de curso, modalidad (virtual/mixto), antiguedad del estudiante y canal de registro.")
heading(d, 1, "Encuestas Transaccionales")
para(d, "Adicionalmente, se envia una encuesta transaccional despues de cada interaccion con soporte y al completar cada curso. Estas encuestas permiten identificar puntos de dolor de forma temprana.")
heading(d, 1, "Planes de Accion")
para(d, "Si el NPS cae mas de 5 puntos en un trimestre, se activa un plan de accion inmediato: analisis de causas, entrevistas cualitativas a detractores, y plan de mejora con responsable y fecha limite.")
d.save(str(sh / "cal" / "CAL-003.docx"))
print("CAL-003.docx OK")

# ── COM-004 Fidelizacion y Referidos ──
d = docx.Document()
heading(d, 0, "Programa de Fidelizacion y Referidos")
heading(d, 1, "Programa de Referidos")
para(d, "Cada estudiante puede invitar a amigos a traves de un link unico de referido. Cuando el amigo se inscribe en un curso pago, el referidor recibe un credito de $20 USD aplicable a su proxima compra. El amigo recibe un descuento del 10% en su primer curso.")
heading(d, 1, "Beneficios por Lealtad")
para(d, "Despues de completar 5 cursos, el estudiante obtiene el nivel 'Bronce': 5% de descuento en todos los cursos. 10 cursos: nivel 'Plata' con 10% de descuento + acceso anticipado a nuevos cursos. 20 cursos: nivel 'Oro' con 15% de descuento + curso gratuito por ano + certificados sin costo.")
heading(d, 1, "Puntos por Participacion")
para(d, "Los estudiantes ganan puntos por: completar encuestas (50 pts), participar en foros (10 pts por respuesta util), asistir a webinars (100 pts), y dejar resenas de cursos (75 pts). Cada 1000 puntos equivalen a $10 USD de descuento.")
d.save(str(sh / "com" / "COM-004.docx"))
print("COM-004.docx OK")

# ── RH-007 Capacitacion y Desarrollo ──
d = docx.Document()
heading(d, 0, "Politica de Capacitacion y Desarrollo")
heading(d, 1, "Presupuesto de Formacion")
para(d, "Cada colaborador tiene un presupuesto anual de $1,500 USD para formacion externa (cursos, certificaciones, conferencias). El presupuesto se renueva en enero y no es acumulable entre anos.")
heading(d, 1, "Plan de Desarrollo Individual")
para(d, "Cada colaborador elabora un Plan de Desarrollo Individual (PDI) al inicio de cada ano, en conjunto con su lider. El PDI identifica brechas de competencias y define acciones de desarrollo para el ano.")
heading(d, 1, "Programa de Mentoria")
para(d, "Los colaboradores pueden solicitar un mentor dentro de la organizacion. El programa tiene una duracion de 6 meses con sesiones quincenales. Los mentores reciben un reconocimiento especial por su participacion.")
heading(d, 1, "Capacitaciones Obligatorias")
para(d, "Todos los colaboradores deben completar anualmente: curso de seguridad de la informacion (2 horas), curso de etica y compliance (1 hora), curso de diversidad e inclusion (1 hora), y curso de proteccion de datos (1 hora).")
d.save(str(sh / "rh" / "RH-007.docx"))
print("RH-007.docx OK")

# ── RH-008 Igualdad y No Discriminacion ──
d = docx.Document()
heading(d, 0, "Politica de Igualdad y No Discriminacion")
heading(d, 1, "Compromiso")
para(d, "AprendeYa se compromete a mantener un entorno de trabajo libre de discriminacion por raza, color, religion, genero, orientacion sexual, identidad de genero, edad, discapacidad, origen nacional o cualquier otra caracteristica protegida por la ley.")
heading(d, 1, "Igualdad Salarial")
para(d, "AprendeYa realiza auditorias salariales anuales para garantizar que no existan brechas injustificadas. Los salarios se determinan exclusivamente por: rol, experiencia, desempeno y responsabilidades.")
heading(d, 1, "Seleccion y Promocion")
para(d, "Los procesos de seleccion y promocion se basan en meritocracia. Se utilizan paneles diversos y criterios objetivos. Se fomenta la postulacion de grupos subrepresentados en posiciones de liderazgo.")
heading(d, 1, "Ajustes Razonables")
para(d, "AprendeYa proporciona ajustes razonables para colaboradores con discapacidades: horarios flexibles, equipos adaptados, espacios accesibles y software de asistencia.")
heading(d, 1, "Canal de Denuncia")
para(d, "Cualquier acto de discriminacion puede ser reportado de forma confidencial a inclusion@aprendeya.com o a traves de la Linea Etica sin temor a represalias.")
d.save(str(sh / "rh" / "RH-008.docx"))
print("RH-008.docx OK")

# ── OPE-005 Manual de Infraestructura Fisica ──
d = docx.Document()
heading(d, 0, "Manual de Infraestructura Fisica")
heading(d, 1, "Oficinas")
para(d, "AprendeYa opera desde tres sedes: Bogota (sede principal, Calle 100 # 15-45, 4 pisos), Medellin (oficina satelite, Centro Empresarial San Fernando, piso 8) y Cali (coworking WeWork, Avenida 6N # 25-30).")
heading(d, 1, "Horarios y Acceso")
para(d, "Horario de oficinas: lunes a viernes 7am a 7pm. Acceso con tarjeta de proximidad individual. Registro de ingreso y salida obligatorio. Visitas deben registrarse en recepcion con 24 horas de anticipacion.")
heading(d, 1, "Equipamiento")
para(d, "Cada colaborador recibe: laptop corporativa (ThinkPad o MacBook Pro), monitor externo, teclado y mouse ergonomicos, audifonos con cancelacion de ruido, escritorio ajustable de altura. El inventario se actualiza cada 3 anos.")
heading(d, 1, "Seguridad Fisica")
para(d, "Camaras de vigilancia en areas comunes. Extintores revisados trimestralmente. Salidas de emergencia senalizadas. Simulacros de evacuacion semestrales. Botiquin de primeros auxilios en cada piso.")
d.save(str(sh / "ope" / "OPE-005.docx"))
print("OPE-005.docx OK")

# ── FIN-003 Estados Financieros 2025 ──
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Balance General 2025"
ws.append(["Cuenta", "Dic 2025", "Dic 2024"])
for row in [
    ["Activo Corriente", 850000000, 620000000],
    ["Efectivo", 320000000, 210000000],
    ["Cuentas por Cobrar", 380000000, 290000000],
    ["Inventarios", 150000000, 120000000],
    ["Activo No Corriente", 1200000000, 950000000],
    ["Propiedad y Equipo", 780000000, 610000000],
    ["Intangibles", 420000000, 340000000],
    ["TOTAL ACTIVO", 2050000000, 1570000000],
    ["", "", ""],
    ["Pasivo Corriente", 520000000, 410000000],
    ["Cuentas por Pagar", 280000000, 230000000],
    ["Obligaciones Laborales", 180000000, 140000000],
    ["Impuestos por Pagar", 60000000, 40000000],
    ["Pasivo No Corriente", 380000000, 290000000],
    ["TOTAL PASIVO", 900000000, 700000000],
    ["", "", ""],
    ["Patrimonio", 1150000000, 870000000],
    ["Capital Social", 500000000, 500000000],
    ["Utilidades Retenidas", 650000000, 370000000],
    ["TOTAL PATRIMONIO", 1150000000, 870000000],
]:
    ws.append(row)
ws2 = wb.create_sheet("P&L 2025")
ws2.append(["Cuenta", "2025", "2024"])
for row in [
    ["Ingresos Operacionales", 3200000000, 2100000000],
    ["Matriculas", 1800000000, 1200000000],
    ["Suscripciones", 950000000, 600000000],
    ["Cursos Corporativos", 450000000, 300000000],
    ["Costo de Ventas", -1400000000, -950000000],
    ["Utilidad Bruta", 1800000000, 1150000000],
    ["Gastos Operativos", -1200000000, -850000000],
    ["Tecnologia", -350000000, -250000000],
    ["Marketing", -280000000, -200000000],
    ["Nomina", -480000000, -320000000],
    ["Generales", -90000000, -80000000],
    ["Utilidad Operativa", 600000000, 300000000],
    ["Gastos Financieros", -40000000, -30000000],
    ["Utilidad Neta", 560000000, 270000000],
]:
    ws2.append(row)
wb.save(str(sh / "fin" / "FIN-003.xlsx"))
print("FIN-003.xlsx OK")

# ── FIN-004 Proyeccion Financiera ──
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Proyeccion 2026-2027"
ws.append(["Indicador", "2025 Real", "2026 Proy", "2027 Proy"])
for row in [
    ["Ingresos", 3200000000, 4500000000, 5800000000],
    ["Crecimiento (%)", "", 41, 29],
    ["Estudiantes Pagos", 8500, 15000, 22000],
    ["Ticket Promedio", 376000, 300000, 264000],
    ["CAC (USD)", 25, 20, 18],
    ["LTV (USD)", 150, 200, 250],
    ["Margen Bruto (%)", 56, 60, 62],
    ["Margen Neto (%)", 18, 22, 25],
    ["NPS", 45, 55, 65],
    ["Colaboradores", 45, 65, 85],
    ["Cursos en Catalogo", 40, 80, 120],
]:
    ws.append(row)
ws2 = wb.create_sheet("Supuestos")
ws2.append(["Variable", "Valor"])
ws2.append(["Tasa de conversion lead-estudiante", "4.5%"])
ws2.append(["Tasa de retencion mensual", "88%"])
ws2.append(["Renovacion anual planes premium", "75%"])
ws2.append(["Crecimiento catalogo anual", "100%"])
ws2.append(["Inversion en tecnologia (% ingresos)", "12%"])
ws2.append(["Inversion en marketing (% ingresos)", "10%"])
wb.save(str(sh / "fin" / "FIN-004.xlsx"))
print("FIN-004.xlsx OK")

# ── FIN-005 Reembolso de Gastos ──
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Politica"
ws.append(["Concepto", "Tope Diario", "Requiere Autorizacion", "Soporte"])
ws.append(["Alimentacion", 45000, "No", "Factura"])
ws.append(["Transporte terrestre", 80000, "No", "Factura / tiquete"])
ws.append(["Transporte aereo nacional", "", "Si - Lider", "Tiquete + boarding"])
ws.append(["Hospedaje", 250000, "Si - Lider", "Factura"])
ws.append(["Viaticos internacionales", "USD 120/dia", "Si - Direccion", "Factura + itinerario"])
ws.append(["Gasolina (vehiculo propio)", "$1,800/km", "No", "Registro de kilometraje"])
ws.append(["Peajes y parqueaderos", "Costo real", "No", "Factura"])
ws.append(["Papeleria y suministros", 50000, "No", "Factura"])
ws.append(["Suscripciones herramientas", "", "Si - TI", "Factura"])
ws2 = wb.create_sheet("Procedimiento")
ws2.append(["Paso", "Descripcion", "Plazo"])
ws2.append(["1", "Realizar el gasto con recursos propios", ""])
ws2.append(["2", "Solicitar factura a nombre de AprendeYa (NIT 901.XXX.XXX)", ""])
ws2.append(["3", "Ingresar al sistema de gastos y subir la factura", "Dentro de 15 dias del gasto"])
ws2.append(["4", "El lider directo aprueba el gasto", "3 dias habiles"])
ws2.append(["5", "Finanzas procesa el reembolso", "5 dias habiles"])
ws2.append(["6", "El valor se acredita en la nomina siguiente", ""])
wb.save(str(sh / "fin" / "FIN-005.xlsx"))
print("FIN-005.xlsx OK")

print("\nTODOS LOS DOCUMENTOS BINARIOS GENERADOS")
