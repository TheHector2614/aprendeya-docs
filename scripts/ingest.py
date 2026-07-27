"""
Ingesta documental — AprendeYa v2
Procesa los documentos listados en docs-management/inventario.yaml:
  1. Extrae texto plano según el formato (HTML, DOCX, XLSX, PPTX, PDF, MD, CSV, JSON, TXT)
  2. Limpia el texto (elimina ruido, encabezados, pies de página)
  3. Divide en chunks (estructural por secciones, con fallback por tamaño fijo)
  4. Genera embeddings (modelo multilingüe)
  5. Indexa en ChromaDB con metadatos completos

Uso: python scripts/ingest.py
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import yaml
from pipeline import Config, Extractor, Cleaner, Chunker, Embedder, Indexer


def load_inventario(path: Path) -> list[dict]:
    with open(path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return data.get("documentos", [])


def resolve_path(doc: dict) -> Path | None:
    fuente: str = doc.get("fuente", "")
    formato: str = doc.get("formato", "HTML").lower()
    ext_map = {
        "html": ".html", "md": ".md", "docx": ".docx",
        "xlsx": ".xlsx", "pptx": ".pptx", "pdf": ".pdf",
        "txt": ".txt", "csv": ".csv", "json": ".json",
    }
    ext = ext_map.get(formato, ".html")

    if "Sitio web" in fuente:
        ruta = Config.RAW_DIR / "web" / f"{doc['id']}.html"
    elif "Repositorio" in fuente:
        ruta = Config.RAW_DIR / "github" / f"{doc['id']}.md"
    elif "Google Drive" in fuente:
        ruta = Config.RAW_DIR / "drive" / doc["categoria"].lower() / f"{doc['id']}{ext}"
    elif "SharePoint" in fuente:
        ruta = Config.RAW_DIR / "sharepoint" / doc["categoria"].lower() / f"{doc['id']}{ext}"
    else:
        ruta = Config.RAW_DIR / "general" / f"{doc['id']}{ext}"

    if formato == "html" and ruta.suffix != ".html":
        ruta = ruta.with_suffix(".html")
    return ruta


def ensure_sample_docs():
    web_dir = Config.RAW_DIR / "web"
    web_dir.mkdir(parents=True, exist_ok=True)

    pages_dir = Config.ROOT / "src" / "pages"
    doc_map = {
        "reglamento-estudiante": "ACA-001",
        "politica-reembolso": "ACA-002",
        "faq": "ACA-003",
        "guia-uso": "ACA-004",
        "programa-becas": "ACA-005",
    }

    for stem, doc_id in doc_map.items():
        out = web_dir / f"{doc_id}.html"
        if not out.exists():
            astro_file = pages_dir / f"{stem}.astro"
            if astro_file.exists():
                content = astro_file.read_text("utf-8")
                simple_html = f"<html><body><pre>{content}</pre></body></html>"
                out.write_text(simple_html, "utf-8")
                print(f"  [creado] {out.name}")

    drive_rh = Config.RAW_DIR / "drive" / "rh"
    drive_rh.mkdir(parents=True, exist_ok=True)
    showplaces = {
        drive_rh / "RH-001.docx": "Manual de inducción para empleados de AprendeYa.",
        drive_rh / "RH-002.docx": "Política de trabajo remoto: horarios flexibles, conectividad, ergonomía.",
    }
    for path, content in showplaces.items():
        if not path.exists():
            from docx import Document
            doc = Document()
            doc.add_heading(path.stem.replace("-", " "), level=1)
            for para in content.split("\n"):
                doc.add_paragraph(para.strip())
            doc.save(str(path))
            print(f"  [creado] {path.name}")


def main():
    cfg = Config()
    extractor = Extractor(ocr_fallback=cfg.OCR_FALLBACK)
    cleaner = Cleaner()
    chunker = Chunker(
        strategy=cfg.CHUNK_STRATEGY,
        size=cfg.CHUNK_SIZE,
        overlap=cfg.CHUNK_OVERLAP,
    )
    embedder = Embedder(model_name=cfg.EMBEDDING_MODEL)
    indexer = Indexer(
        persist_dir=str(cfg.INDEX_DIR),
        collection_name=cfg.COLLECTION_NAME,
        reset=True,
    )

    print(f"=== INGESTA DOCUMENTAL v2 — AprendeYa ===")
    print(f"Estrategia: {cfg.CHUNK_STRATEGY} | Chunk size: {cfg.CHUNK_SIZE}")
    print()

    inventario = load_inventario(cfg.INVENTARIO)
    print(f"Documentos en inventario: {len(inventario)}")

    ensure_sample_docs()

    all_chunks = []
    for doc in inventario:
        doc_id = doc["id"]
        titulo = doc.get("titulo", doc_id)
        categoria = doc.get("categoria", "GENERAL")

        ruta = resolve_path(doc)
        if not ruta or not ruta.exists():
            print(f"  [omitido] {doc_id} — archivo no encontrado: {ruta}")
            continue

        try:
            text = extractor.extract(ruta)
            text = cleaner.clean(text)
            chunks = chunker.chunk(text, doc_id, titulo, categoria)
            all_chunks.extend(chunks)
            secciones = set(c.seccion for c in chunks if c.seccion)
            secciones_str = f", secciones: {', '.join(secciones)}" if secciones else ""
            print(f"  [ok] {doc_id} — {titulo} ({len(chunks)} chunks{secciones_str})")
        except ValueError as e:
            print(f"  [error] {doc_id} — {e}")
        except Exception as e:
            print(f"  [error] {doc_id} — {e}")

    print(f"\nTotal chunks generados: {len(all_chunks)}")

    if not all_chunks:
        print("No hay chunks para indexar.")
        return

    texts = [c.text for c in all_chunks]
    print("Generando embeddings...")
    embeddings = embedder.embed(texts, show_progress=True)

    print("Indexando en ChromaDB...")
    indexer.index(all_chunks, embeddings)

    print(f"\n=== INGESTA COMPLETADA ===")
    print(f"Chunks indexados: {indexer.count()}")
    print(f"Colección: {cfg.COLLECTION_NAME}")
    print(f"Índice: {cfg.INDEX_DIR}")


if __name__ == "__main__":
    main()
