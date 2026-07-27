import docx
from pathlib import Path

BASE = Path("raw")
gh = BASE / "github"
sh = BASE / "sharepoint"

# ════════════════════════════════════════════════
# TEC-003.md — Manual de APIs y Webhooks (expandido)
# ════════════════════════════════════════════════
md = """# Manual de APIs y Webhooks

## Autenticacion y Seguridad
Todas las llamadas a la API de AprendeYa requieren autenticacion mediante una API Key que debe enviarse en el header `X-API-Key`. Las API Keys se generan desde el panel de desarrollador en la plataforma. Cada key tiene permisos asociados al rol del usuario que la crea. Las keys de instructores solo pueden acceder a endpoints de cursos propios. Las keys de administradores tienen acceso completo.

## Endpoints de Cursos
El endpoint `GET /api/v1/cursos` permite listar todos los cursos con paginacion de 20 resultados por pagina. Soporta filtros por categoria, nivel y precio. `GET /api/v1/cursos/{id}` devuelve el detalle completo de un curso incluyendo modulos, lecciones y requisitos. `POST /api/v1/cursos` crea un nuevo curso y requiere rol de instructor. `PUT /api/v1/cursos/{id}` actualiza un curso existente. `DELETE /api/v1/cursos/{id}` elimina un curso y solo el creador o un admin pueden hacerlo.

## Endpoints de Estudiantes
`GET /api/v1/estudiantes` lista los estudiantes inscritos a tus cursos con informacion basica de perfil. `GET /api/v1/estudiantes/{id}/progreso` muestra el avance del estudiante en cada modulo, incluyendo lecciones completadas, calificaciones de cuestionarios y tiempo invertido. `POST /api/v1/estudiantes/{id}/certificado` genera el certificado de finalizacion cuando el estudiante completa el curso con nota minima de 60/100.

## Endpoints de Pagos
`GET /api/v1/pagos` permite consultar el historial de transacciones de un estudiante, incluyendo fecha, monto, metodo de pago y estado. `POST /api/v1/pagos/reembolso` inicia el proceso de reembolso de una matricula. Requiere especificar el motivo y el numero de transaccion original. El reembolso se procesa en un maximo de 5 dias habiles.

## Rate Limiting
La API tiene un limite de 1000 requests por minuto por API Key. Si se excede este limite, la API responde con codigo 429 (Too Many Requests) y un mensaje indicando cuando se restablecera el limite. Cada respuesta incluye headers con el estado actual del rate limit: `X-RateLimit-Limit`, `X-RateLimit-Remaining` y `X-RateLimit-Reset`.

## Webhooks Disponibles
La plataforma expone webhooks para eventos en tiempo real. `curso.creado` se dispara cuando un instructor publica un nuevo curso. `curso.actualizado` cuando se modifica el contenido. `estudiante.inscrito` cuando un estudiante se inscribe a un curso. `estudiante.completado` al finalizar exitosamente. `pago.exitosa` cuando se procesa un pago. `pago.reembolsado` cuando se completa un reembolso.

## Configuracion de Webhooks
Para configurar un webhook, registra tu URL en el panel de desarrollador. La plataforma enviara un POST con el payload del evento en formato JSON. Tu servidor debe responder con 200 OK dentro de 5 segundos. Si no recibe respuesta, la plataforma reintenta hasta 3 veces con backoff exponencial (5s, 30s, 120s). Puedes configurar hasta 5 URLs de webhook diferentes.
"""
Path(gh / "TEC-003.md").write_text(md, encoding="utf-8")
print("TEC-003.md OK (expandido)")

# ════════════════════════════════════════════════
# TEC-004.md — Politica de Seguridad Informatica (expandido)
# ════════════════════════════════════════════════
md = """# Politica de Seguridad Informatica

## Controles de Acceso
La autenticacion multifactor (MFA) es obligatoria para todos los colaboradores de AprendeYa, incluyendo personal administrativo, instructores y contratistas. Se aplica el principio de minimo privilegio: cada usuario tiene acceso unicamente a los recursos necesarios para su funcion. Los accesos y permisos se revisan cada trimestre. Las cuentas de servicio rotan sus contrasenas cada 90 dias. El acceso remoto a la red interna solo es posible a traves de VPN corporativa.

## Gestion de Parches y Vulnerabilidades
Los parches de seguridad criticos deben aplicarse en menos de 48 horas desde su publicacion. Los parches de alta prioridad tienen un plazo de 7 dias. Los parches de prioridad media o baja se aplican en el ciclo mensual de mantenimiento. Se realiza un escaneo de vulnerabilidades semanal automatizado con Trivy para contenedores y Snyk para dependencias. Los resultados se revisan en el comite de seguridad quincenal.

## Cifrado y Proteccion de Datos
Todas las bases de datos utilizan cifrado AES-256 en reposo. Las comunicaciones entre servicios usan TLS 1.3. Los backups se realizan diariamente con retencion de 30 dias y se almacenan en una region separada. Se realizan pruebas de restauracion mensuales para verificar la integridad de los backups. Los datos de tarjetas de credito nunca se almacenan en la plataforma; son procesados directamente por Stripe.

## Seguridad de Red
La infraestructura cuenta con firewalls de aplicacion web (WAF) en todos los endpoints publicos, con reglas basadas en OWASP Top 10. La red esta segmentada por ambiente: desarrollo, staging y produccion, con aislamiento completo entre ellos. Todo el trafico interno entre microservicios esta cifrado y autenticado mediante mTLS. El monitoreo de trafico incluye deteccion de intrusiones (IDS) con alertas en tiempo real.

## Dispositivos y Equipos
Todos los dispositivos corporativos estan administrados via MDM con cifrado de disco completo. La politica de bloqueo de pantalla exige bloqueo automatico tras 5 minutos de inactividad. Esta prohibido conectar dispositivos de almacenamiento USB no autorizados. Todos los equipos deben tener antivirus actualizado con escaneo en tiempo real. Los dispositivos personales no pueden acceder a recursos corporativos.
"""
Path(gh / "TEC-004.md").write_text(md, encoding="utf-8")
print("TEC-004.md OK (expandido)")

# ════════════════════════════════════════════════
# TEC-005.md — Guia de Estandares de Codigo (expandido)
# ════════════════════════════════════════════════
md = """# Guia de Estandares de Codigo

## Convenciones Generales de Lenguaje
En AprendeYa usamos Python 3.12+ para el backend de servicios (FastAPI), TypeScript 5.x para el frontend (Next.js y React), y Java 21 para el servicio de pagos (Spring Boot). El estilo de codigo se verifica con Ruff para Python, ESLint con configuracion estandar para TypeScript, y Checkstyle para Java. El formato automatico se aplica con Ruff format en Python y Prettier en TypeScript. La tabulacion es de 4 espacios para Python y 2 espacios para TypeScript. Los archivos deben terminar con una linea en blanco y no tener espacios al final de las lineas.

## Estrategia de Branching (Git Flow)
Todas las ramas se basan en la siguiente convencion: `main` es la rama de produccion y solo recibe merges via Pull Request aprobado. `develop` es la rama de integracion donde confluyen las funcionalidades antes de pasar a produccion. `feature/{nombre}` se crea a partir de develop para nuevas funcionalidades. `fix/{nombre}` se crea a partir de develop para correcciones. `hotfix/{nombre}` se crea a partir de main para correcciones urgentes en produccion. Los hotfixes deben mergearse tanto a main como a develop.

## Proceso de Code Review
Todo cambio en el codigo debe pasar por un Pull Request con al menos una aprobacion de un par. El PR debe incluir tests automatizados que pasen exitosamente en CI. Esta prohibido mergear el propio PR sin aprobacion externa. El maximo recomendado es de 400 lineas por PR; cambios mayores deben justificarse y dividirse en PRs mas pequenos. Los comentarios del revisor deben resolverse antes del merge. Las revisiones deben completarse en menos de 24 horas habiles.

## Formato de Commits
Los commits siguen el formato conventional commits con los prefijos: feat (nueva funcionalidad), fix (correccion de bug), chore (tareas de mantenimiento), docs (documentacion), refactor (refactorizacion sin cambio funcional), test (agregar o modificar tests), style (cambios de formato). El formato es `tipo(scope): mensaje en ingles imperativo`. Ejemplo: `feat(courses): add search by category endpoint`. El mensaje debe ser descriptivo pero conciso, maximo 72 caracteres en la primera linea.

## Cobertura y Tipos de Tests
La cobertura minima de tests unitarios es del 80% y de tests de integracion del 60%. Los tests unitarios son obligatorios para toda logica de negocio, incluyendo validaciones, calculos y transformaciones de datos. Los tests de integracion deben cubrir todos los endpoints de API. Los tests E2E son obligatorios unicamente para flujos criticos como login, checkout y generacion de certificados. Los tests se ejecutan en cada PR y en el pipeline de CI antes del despliegue a produccion.
"""
Path(gh / "TEC-005.md").write_text(md, encoding="utf-8")
print("TEC-005.md OK (expandido)")

# ════════════════════════════════════════════════
# CAL-003.docx — Politica de Satisfaccion NPS (mejorado)
# ════════════════════════════════════════════════
d = docx.Document()
d.add_heading("Politica de Satisfaccion del Cliente (NPS)", 0)
d.add_heading("Que es el NPS", 1)
d.add_paragraph("El Net Promoter Score (NPS) es la metrica principal de satisfaccion de AprendeYa. Se calcula preguntando a los estudiantes: 'En una escala del 0 al 10, que tan probable es que recomiendes AprendeYa a un amigo o colega?' Los estudiantes que responden 9 o 10 son 'Promotores'. Los que responden 7 u 8 son 'Pasivos'. Los que responden 0 a 6 son 'Detractores'. El NPS se calcula restando el porcentaje de Detractores del porcentaje de Promotores.")
d.add_heading("Metas de NPS", 1)
d.add_paragraph("El NPS actual de AprendeYa es 45. La meta para 2026 es alcanzar 55 puntos. La meta para 2027 es llegar a 65 puntos. El NPS se segmenta por tipo de curso (idiomas, programacion, negocios), por modalidad (virtual vs mixto), por antiguedad del estudiante y por canal de registro. Cada segmento tiene metas especificas definidas en el plan anual.")
d.add_heading("Encuestas Transaccionales", 1)
d.add_paragraph("Adicionalmente al NPS trimestral, se envian encuestas transaccionales cortas despues de cada interaccion con soporte y al completar cada curso. Estas encuestas permiten detectar puntos de dolor de forma temprana y tomar acciones correctivas antes de la medicion trimestral. La tasa de respuesta esperada es superior al 30%.")
d.add_heading("Planes de Accion Correctivos", 1)
d.add_paragraph("Si el NPS cae mas de 5 puntos en un trimestre, se activa un plan de accion inmediato que incluye: analisis de causas raiz, entrevistas cualitativas con al menos 10 detractores, identificacion de los principales puntos de dolor, y un plan de mejora con responsable asignado y fecha limite. El plan se revisa en el comite directivo mensual.")
d.add_heading("Reportes y Seguimiento", 1)
d.add_paragraph("El NPS se reporta mensualmente al comite directivo y trimestralmente a toda la organizacion. Los resultados se publican en el dashboard de calidad accesible para todos los colaboradores. Cada equipo tiene visibilidad del NPS de los cursos o servicios bajo su responsabilidad.")
d.save(str(sh / "cal" / "CAL-003.docx"))
print("CAL-003.docx OK (mejorado)")

# ════════════════════════════════════════════════
# COM-003.docx — Convenios Corporativos (mejorado)
# ════════════════════════════════════════════════
d = docx.Document()
d.add_heading("Programa de Convenios Corporativos", 0)
d.add_heading("Beneficios por Volumen", 1)
d.add_paragraph("Las empresas aliadas de AprendeYa acceden a descuentos progresivos segun el numero de empleados inscritos. Para grupos de 10 a 49 empleados, el descuento es del 15% sobre el precio del catalogo. Para grupos de 50 a 199 empleados, el descuento es del 25%. Para grupos de 200 o mas empleados, el descuento alcanza el 35% e incluye reportes de progreso personalizados con dashboards ejecutivos.")
d.add_heading("Requisitos para Empresas", 1)
d.add_paragraph("La empresa debe firmar un convenio marco con vigencia minima de un ano. El pago puede ser mensual, trimestral o anual, segun prefiera la empresa. Las organizaciones pueden elegir entre el catalogo completo de cursos (modelo abierto) o una seleccion personalizada de cursos alineados con sus necesidades de capacitacion.")
d.add_heading("Facturacion y Gestion", 1)
d.add_paragraph("La facturacion corporativa se realiza a 30 dias con orden de compra. Las empresas pueden agregar o retirar empleados de forma mensual sin penalizacion. El reporte de uso y progreso se entrega el primer dia habil de cada mes e incluye: cursos iniciados, cursos completados, horas de capacitacion y certificaciones obtenidas.")
d.add_heading("Contacto Comercial", 1)
d.add_paragraph("Las empresas interesadas pueden contactar al equipo corporativo a traves del correo corporativo@aprendeya.com o llamando al 01-800-123-4567 opcion 2. Un ejecutivo comercial dedicado acompanara todo el proceso, desde la cotizacion hasta la implementacion y el seguimiento mensual.")
d.save(str(sh / "com" / "COM-003.docx"))
print("COM-003.docx OK (mejorado)")

# ════════════════════════════════════════════════
# COM-004.docx — Fidelizacion y Referidos (mejorado)
# ════════════════════════════════════════════════
d = docx.Document()
d.add_heading("Programa de Fidelizacion y Referidos", 0)
d.add_heading("Programa de Referidos", 1)
d.add_paragraph("Cada estudiante de AprendeYa recibe un link unico de referido que puede compartir con amigos y familiares. Cuando una persona se inscribe en un curso pago usando ese link, el estudiante que refirio recibe un credito de $20 USD que puede aplicar a su proxima compra. La persona referida recibe un descuento del 10% en su primer curso. No hay limite en la cantidad de referidos que un estudiante puede generar.")
d.add_heading("Niveles de Fidelizacion", 1)
d.add_paragraph("El programa de fidelidad tiene tres niveles basados en la cantidad de cursos completados. Nivel Bronce: al completar 5 cursos, el estudiante obtiene un 5% de descuento permanente en todos los cursos del catalogo. Nivel Plata: al completar 10 cursos, el descuento aumenta al 10% y se obtiene acceso anticipado a los nuevos cursos lanzados. Nivel Oro: al completar 20 cursos, el descuento es del 15%, se recibe un curso gratuito por ano y los certificados no tienen costo adicional.")
d.add_heading("Puntos por Participacion", 1)
d.add_paragraph("Ademas de los descuentos por nivel, los estudiantes acumulan puntos canjeables por descuentos. Se ganan 50 puntos por completar encuestas de satisfaccion, 10 puntos por cada respuesta util en los foros (marcada por el instructor), 100 puntos por asistir a webinars en vivo, y 75 puntos por dejar resenas de cursos. Cada 1000 puntos equivalen a $10 USD de descuento aplicable en cualquier curso.")
d.save(str(sh / "com" / "COM-004.docx"))
print("COM-004.docx OK (mejorado)")

print("\nTodos los documentos mejorados OK")
